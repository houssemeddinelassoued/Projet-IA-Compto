from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Créer le document
doc = Document()

# Configurer les marges
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Style par défaut
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# --- EN-TÊTE ENTREPRISE ---
header = doc.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = header.add_run("ENTREPRISE ÉNERGIE SOLAIRE")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 100, 0)

header.add_run("\n")

info_lines = [
    "Zone Industrielle, Raf Raf",
    "7024 Bizerte, Tunisie",
    "Tél : +216 72 000 000",
    "Email : contact@entreprise-solaire.tn",
    "MF : 1234567/A"
]

for line in info_lines:
    run = header.add_run(line + "\n")
    run.font.size = Pt(10)

# Ligne de séparation
doc.add_paragraph("─" * 70)

# --- TITRE DEVIS ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("DEVIS D'INSTALLATION PHOTOVOLTAÏQUE")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 80, 0)

doc.add_paragraph()

# --- INFOS DEVIS ET CLIENT ---
info_table = doc.add_table(rows=1, cols=2)
info_table.autofit = True

cell_left = info_table.rows[0].cells[0]
cell_right = info_table.rows[0].cells[1]

# Informations devis
p = cell_left.paragraphs[0]
p.add_run("N° Devis : ").bold = True
p.add_run("2026-01-001\n")
p.add_run("Date : ").bold = True
p.add_run("24 Janvier 2026\n")
p.add_run("Validité de l'offre : ").bold = True
p.add_run("30 jours")

# Informations client
p = cell_right.paragraphs[0]
run = p.add_run("CLIENT :\n")
run.bold = True
p.add_run("M. [Nom du Client]\n")
p.add_run("Adresse du chantier\n")
p.add_run("[Code Postal] [Ville]")

doc.add_paragraph()

# --- SECTION DÉTAIL DES PRESTATIONS ---
section_title = doc.add_paragraph()
run = section_title.add_run("Détail des prestations")
run.bold = True
run.font.size = Pt(14)

# Tableau des prestations
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# En-têtes
headers = ["Désignation", "Qté", "P.U. (TND)", "Total (TND)"]
header_row = table.rows[0]
for i, header_text in enumerate(headers):
    cell = header_row.cells[i]
    cell.text = header_text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Fond gris pour l'en-tête
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
    cell._tc.get_or_add_tcPr().append(shading)

# Données
data = [
    ("Onduleur Photovoltaïque Hybride 5kW\n(Marque Premium, MPPT intégré, Garantie 10 ans)", "1", "3 500,00", "3 500,00"),
    ("Panneaux Solaires Monocristallins 400Wc\n(Technologie Half-Cut)", "10", "550,00", "5 500,00"),
    ("Coffret de protection AC/DC\n(Disjoncteurs, Parafoudre, Interrupteur sectionneur)", "1", "1 100,00", "1 100,00"),
    ("Câblage et structure de fixation\n(Rails aluminium, câbles solaires 6mm²)", "1", "1 200,00", "1 200,00"),
    ("Main d'œuvre\n(Pose, Installation, Mise en service & Test)", "1", "2 500,00", "2 500,00"),
]

for row_idx, row_data in enumerate(data, start=1):
    row = table.rows[row_idx]
    for col_idx, value in enumerate(row_data):
        cell = row.cells[col_idx]
        cell.text = value
        if col_idx == 0:
            # Première colonne alignée à gauche
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Autres colonnes centrées
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Ajuster largeur des colonnes
widths = [Cm(9), Cm(2), Cm(3), Cm(3.5)]
for row in table.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width

doc.add_paragraph()

# --- RÉCAPITULATIF FINANCIER ---
section_title = doc.add_paragraph()
run = section_title.add_run("Récapitulatif Financier")
run.bold = True
run.font.size = Pt(14)

recap_table = doc.add_table(rows=3, cols=2)
recap_table.style = 'Table Grid'

recap_data = [
    ("Total Hors Taxes (HT) :", "13 800,00 TND"),
    ("TVA (19%) :", "2 622,00 TND"),
    ("TOTAL TTC À PAYER :", "16 422,00 TND"),
]

for row_idx, (label, value) in enumerate(recap_data):
    row = recap_table.rows[row_idx]
    row.cells[0].text = label
    row.cells[1].text = value
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if row_idx == 0 or row_idx == 2:
        row.cells[1].paragraphs[0].runs[0].bold = True
    if row_idx == 2:
        # Fond jaune pour le total
        for cell in row.cells:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF2CC"/>')
            cell._tc.get_or_add_tcPr().append(shading)

# Ajuster largeur
for row in recap_table.rows:
    row.cells[0].width = Cm(8)
    row.cells[1].width = Cm(5)

doc.add_paragraph()

# Ligne de séparation
doc.add_paragraph("─" * 70)

# --- CONDITIONS DE RÈGLEMENT ---
cond_title = doc.add_paragraph()
run = cond_title.add_run("Conditions de règlement :")
run.bold = True

conditions = doc.add_paragraph()
conditions.add_run("• 30% d'acompte à la signature du devis.\n")
conditions.add_run("• Solde à la réception de l'installation.")

# --- GARANTIES ---
garantie_title = doc.add_paragraph()
run = garantie_title.add_run("Garanties :")
run.bold = True

garanties = doc.add_paragraph()
garanties.add_run("• Matériel : Selon constructeur (10 à 25 ans).\n")
garanties.add_run("• Installation : Garantie décennale installateur.")

doc.add_paragraph()

# --- SIGNATURE ---
signature = doc.add_paragraph()
signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = signature.add_run("Bon pour accord")
run.bold = True
run.font.size = Pt(12)

signature2 = doc.add_paragraph()
signature2.alignment = WD_ALIGN_PARAGRAPH.CENTER
signature2.add_run("(Date, Signature et mention \"Lu et approuvé\")")
signature2.runs[0].italic = True

# Sauvegarder
output_path = r"q:\National Quantum\2-PROJETS\Compto-cie\projet-ia-compto\Devis_Installation_Photovoltaique_2026-01-001.docx"
doc.save(output_path)
print(f"Document créé avec succès : {output_path}")
